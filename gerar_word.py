from docx import Document
from docx.shared import Pt

# Definindo a variável `data` no Python
data = {
    "funcao": "Pintor + ruído",
    "exames": {
        "admissional": "Exame Clínico | Audiometria | Dermatológico | TGO / TGP/ GGT | Ureia | Creatinina | Espirometria | RX Tórax PA OIT | Acuidade visual",
        "periodico": "Exame Clínico | Audiometria | Dermatológico | TGO / TGP/ GGT | Ureia | Creatinina | Espirometria | RX Tórax PA OIT | Acuidade visual",
        "retorno": "Exame Clínico",
        "mudanca_risco": "Realizar os exames complementares da periodicidade 'Demissional' e acrescentar os exames correspondentes ao 'admissional' do novo GHE.",
        "demissional": "Exame Clínico | Audiometria | Hemograma completo | Espirometria | Raio X Tórax PA OIT"
    },
    "riscos": {
        "acidentes": "Incêndios ou explosões, respingo de produtos químicos nos olhos, choques ou curto circuito, quedas e escorregões",
        "ergonomicos": "Postura inadequada e lesões por esforço repetitivo",
        "fisicos": "Ruído, radiação não ionizante (sol)",
        "biologicos": "N/A",
        "quimicos": "Solventes, pigmentos, resina, cargas e aditivos"
    }
}

# Criar um novo documento
doc = Document()

# Adicionar título
doc.add_heading('Consulta de Exames por Função', 0)

# Adicionar tabela de exames
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'

# Cabeçalho da tabela
header_cells = table.rows[0].cells
header_cells[0].text = 'ADMISSIONAL'
header_cells[1].text = 'PERÍODICO'
header_cells[2].text = 'RETORNO AO TRABALHO'
header_cells[3].text = 'MUDANÇA DE RISCO OCUPACIONAL'
header_cells[4].text = 'DEMISSIONAL'

# Adicionar dados dos exames
exames = data["exames"]
row_cells = table.add_row().cells
row_cells[0].text = exames["admissional"]
row_cells[1].text = exames["periodico"]
row_cells[2].text = exames["retorno"]
row_cells[3].text = exames["mudanca_risco"]
row_cells[4].text = exames["demissional"]

# Adicionar seção de riscos
doc.add_heading('Perigos de Acidentes:', level=1)
doc.add_paragraph(data["riscos"]["acidentes"])

doc.add_heading('Perigos Ergonômicos:', level=1)
doc.add_paragraph(data["riscos"]["ergonomicos"])

doc.add_heading('Perigos Físicos:', level=1)
doc.add_paragraph(data["riscos"]["fisicos"])

doc.add_heading('Perigos Biológicos:', level=1)
doc.add_paragraph(data["riscos"]["biologicos"])

doc.add_heading('Perigos Químicos:', level=1)
doc.add_paragraph(data["riscos"]["quimicos"])

# Salvar o documento
doc.save('Consulta_de_Exames_por_Funcao.docx')